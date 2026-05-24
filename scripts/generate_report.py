"""
Generate PAN26 Task 4 Evaluation Report (docx) following Chinese academic template.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from pathlib import Path


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return h


def add_para(doc, text, bold=False, font_size=10.5, alignment=None, first_line_indent=True):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    run.bold = bold
    return p


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    # ========== TITLE ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run('基于查询分段溯源建模的生成式抄袭源检索方法')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(14)
    run.bold = True

    # ========== AUTHORS ==========
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.space_after = Pt(2)
    run = authors.add_run('王毅¹  韩中元¹(通讯作者)')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)

    # Affiliations
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff.paragraph_format.space_after = Pt(2)
    run = aff.add_run('1. [学校名称] [学院名称], [城市] [邮编]')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)

    aff2 = doc.add_paragraph()
    aff2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff2.paragraph_format.space_after = Pt(2)
    run = aff2.add_run('2. 通讯作者，韩中元，[邮箱待填写]')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)

    email = doc.add_paragraph()
    email.alignment = WD_ALIGN_PARAGRAPH.CENTER
    email.paragraph_format.space_after = Pt(10)
    run = email.add_run('[通讯作者邮箱，请填写]')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)

    # ========== ABSTRACT ==========
    add_heading_styled(doc, '摘要', level=1)

    abstract_text = (
        'PAN@CLEF 2026生成式抄袭检测评测任务（Task 4）旨在从大规模候选语料库中，'
        '为给定的可疑文档检索其原始来源文档。该任务的核心挑战在于：大语言模型生成的可疑文档'
        '通常融合了多个不同来源的内容，将文档作为单一查询进行检索会稀释多源溯源信号。'
        '针对这一问题，本文提出了一种基于查询分段溯源建模（Query-Chunk Provenance Modeling）'
        '的源检索方法。该方法首先将可疑文档按段落结构拆分为若干语义连贯的查询段（query chunk），'
        '每个查询段独立进行BM25检索，获得各自的候选源文档排序列表；'
        '然后通过覆盖率加权投票（coverage-weighted voting）机制聚合各段的检索结果，'
        '奖励被多个查询段同时命中的源文档，构建显式的溯源图（provenance graph）。'
        '在PAN26评测数据集上的实验表明，该方法在所有评价指标上均显著优于标准BM25基线：'
        '在800条查询样本上，R@10从0.920提升至0.970（+5.4%），MRR从0.777提升至0.909（+17.0%）；'
        '在1200条查询样本上，R@10从0.883提升至0.969（+9.8%），MRR从0.753提升至0.909（+20.6%）。'
        '该方法仅依赖numpy和标准Python库，无需GPU或预训练语言模型，'
        '已成功部署为TIRA Docker提交（团队：fosuai，软件：oscillating-list）。'
    )
    add_para(doc, abstract_text, font_size=10.5)

    # Keywords
    kw_para = doc.add_paragraph()
    kw_para.paragraph_format.first_line_indent = Cm(0.74)
    run = kw_para.add_run('关键词：')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)
    run.bold = True
    run = kw_para.add_run('生成式抄袭检测；源检索；查询分段；溯源建模；BM25；覆盖率投票')
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)

    # ========== 1. INTRODUCTION ==========
    add_heading_styled(doc, '1 引言', level=1)

    add_para(doc, (
        '生成式抄袭检测（Generative Plagiarism Detection）是PAN@CLEF 2026新设立的一项评测任务[1]。'
        '与传统逐字抄袭（copy-paste plagiarism）或近义改写抄袭（paraphrase plagiarism）不同，'
        '生成式抄袭指利用大语言模型（Large Language Model, LLM）自动生成文本，'
        '这些文本可能融合、重组或改写自数十个来源文档，形成"多源合成"式的抄袭。'
        '该任务分为两个阶段：（1）源检索（Source Retrieval）——给定一篇可疑文档，'
        '从大规模候选语料库中检索其可能来源的文档集合；'
        '（2）文本对齐（Text Alignment）——在可疑文档与源文档之间精确定位抄袭段落。'
        '本文聚焦于第一阶段，即源检索任务。'
    ))

    add_para(doc, (
        '现有源检索方法普遍采用"文档级查询"（document-level querying）范式：'
        '将整篇可疑文档作为一个完整的查询输入检索系统（无论是基于词法匹配的BM25，'
        '还是基于语义匹配的稠密检索模型如E5、SPECTER），计算可疑文档与每个候选源文档之间的'
        '单一相关性得分。然而，这一范式存在根本性局限：当可疑文档合成自来源A的引言、'
        '来源B的方法论和来源C的结果分析时，查询信号是多个"来源指纹"的噪声混合。'
        '检索模型需要同时匹配三种不同的来源特征，导致排名信号被稀释，'
        '真正的来源文档可能被噪声淹没而无法进入前10名。'
    ))

    add_para(doc, (
        '针对上述多源信号稀释问题，本文提出了一种基于查询分段溯源建模'
        '（Query-Chunk Provenance Modeling）的源检索方法。核心思想是：'
        '不再将可疑文档视为一个整体查询，而是将其按段落结构显式拆分为若干语义连贯的查询段'
        '（query chunks），每个查询段独立检索候选来源，再通过覆盖率加权投票机制聚合各段结果，'
        '重建"可疑文档各部分分别来自哪里"的溯源图（provenance graph）。'
        '该方法的直观动机是：被多个独立查询段同时命中的源文档，'
        '更有可能是真实的全局来源，而非偶发的局部匹配。'
        '在800和1200条查询样本上的实验表明，该方法在R@10、R@100、nDCG@10和MRR四项指标上'
        '均一致优于标准BM25基线。此外，该方法仅依赖numpy和Python标准库，'
        '无需GPU或预训练语言模型，适合在TIRA评测平台中高效部署。'
    ))

    # ========== 2. RELATED WORK ==========
    add_heading_styled(doc, '2 相关工作', level=1)

    add_para(doc, (
        '源检索是抄袭检测流水线的基础环节。自PAN 2014评测[2]设立源检索子任务以来，'
        '参与者提出了多种方法，大致可分为三类：'
    ))

    add_para(doc, (
        '（1）基于词法匹配的方法。这类方法依赖词n-gram重叠、指纹（fingerprinting）'
        '或BM25等词频统计模型进行检索。BM25[3]作为经典的概率检索模型，'
        '通过词频（TF）和逆文档频率（IDF）对候选文档进行评分，'
        '在文本抄袭检测场景中表现强劲——因为抄袭文本通常保留了原文的关键词汇，'
        '词法匹配能够有效捕捉这种"词汇指纹"。在PAN 2014至PAN 2025历届评测[1,2]中，'
        '基于BM25的检索方法始终是强基线。'
    ))

    add_para(doc, (
        '（2）基于语义匹配的方法。随着预训练语言模型的发展，'
        '参与者开始使用Sentence-BERT[4]、E5[5]等句子编码器将文档映射到稠密向量空间，'
        '通过余弦相似度进行语义级匹配。这类方法对改写程度较高的抄袭文本具有更好的泛化能力，'
        '但计算成本远高于词法匹配。PAN 2025评测报告[1]指出，'
        'Linq-Embed-Mistral在召回率上表现最佳（0.82），但所有参与团队的方法'
        '在跨域泛化（PAN 2015数据集）上均显著下降。'
    ))

    add_para(doc, (
        '（3）基于文档切分的方法。对于长文档检索，将文档切分为固定大小的文本块'
        '（chunks），分别编码后通过最大池化或加权平均进行聚合[6]。'
        '这类方法在开放域问答和段落检索中取得了良好效果，但仍将查询视为单一整体，'
        '未对查询侧的多源结构进行建模。'
    ))

    add_para(doc, (
        '与上述方法不同，本文提出的查询分段溯源建模方法从查询侧入手：'
        '将可疑文档显式拆分为多个独立查询段，每个查询段独立检索，'
        '再通过投票机制聚合构建溯源图。这一思路的关键差异在于：它不再假设可疑文档'
        '来自单一来源，而是显式建模多源抄袭场景下的溯源结构。'
    ))

    # ========== 3. METHOD ==========
    add_heading_styled(doc, '3 方法', level=1)

    add_heading_styled(doc, '3.1 任务形式化定义', level=2)
    add_para(doc, (
        '给定一个源文档语料库C = {d₁, d₂, ..., d_N}（共N篇文档）和一个可疑文档集合Q，'
        '对于每个可疑文档q ∈ Q，源检索任务的目标是从C中检索出一个排序文档列表R(q) = [d_{r₁}, d_{r₂}, ..., d_{r_K}]，'
        '使得与q存在真实抄袭关系的源文档尽可能排在列表的前列。'
        '每个可疑文档q可能与多个源文档存在抄袭关系，但评测时每个查询仅关联一个'
        '主要来源文档d*(q)（由qrels标注）。检索质量通过R@K、nDCG@K和MRR等指标衡量。'
    ))

    add_heading_styled(doc, '3.2 查询分段溯源建模', level=2)
    add_para(doc, (
        '本文提出的查询分段溯源建模方法包含三个核心组件：'
        '（1）段落级查询分段（Paragraph-level Query Segmentation）；'
        '（2）独立分段检索（Independent Segment Retrieval）；'
        '（3）覆盖率加权源投票（Coverage-Weighted Source Voting）。'
        '图1给出了方法的整体架构。'
    ))

    add_para(doc,
        '[图1位置：方法整体架构图。左侧：可疑文档输入 → 段落分段 → 查询段C₁, C₂, ..., Cₙ；'
        '中间：每个查询段独立BM25检索 → 候选源文档排序列表；'
        '右侧：覆盖率加权投票聚合 → 最终源文档排序列表。'
        '重点框标注：段落分段模块、覆盖率加权投票模块。]',
        font_size=9)

    add_para(doc, (
        '段落级查询分段。给定可疑文档q，首先按双换行符（\\n\\n）将其拆分为原始段落。'
        '然后采用合并-截断策略：依次遍历段落并将相邻段落合并，直到当前缓冲区长度'
        '超过max_chars（3,000字符），此时将缓冲区内容作为一个独立查询段输出，'
        '并开始新的缓冲区。丢弃长度不足min_chars（800字符）的尾部残段。'
        '每个查询的上限为20个查询段。'
        '该策略产生的查询段大致对应学术文档的章节结构（引言段落、方法描述、结果讨论等），'
        '具有较好的语义连贯性。对于过短或缺乏段落结构的可疑文档，回退为整文档查询。',
    ))

    add_para(doc, (
        '独立分段检索。每个查询段cᵢ独立执行BM25检索，获得其top-50候选源文档列表。'
        '为提升分段检索效率，采用以下优化策略：（a）查询词筛选——仅选择文档频率'
        '不超过5,000（约占语料库的8%）的查询词参与评分，过滤掉区分度低且发布列表极长的'
        '高频词；（b）限制查询词数量——每个查询段最多使用20个查询词（而非标准检索的100个），'
        '因为短文本中的独特判别词较少；（c）高效top-k提取——使用堆排序（heapq.nlargest）'
        '而非全排序，将复杂度从O(N log N)降低至O(N log k)。',
    ))

    add_para(doc, (
        '覆盖率加权源投票。将n个查询段的检索结果聚合为最终排序。'
        '对于每个候选源文档d，其最终得分的计算公式为：'
    ))

    add_para(doc,
        'Score(d) = Σᵢ BM25(d|cᵢ) × (1 + count(d) / n)',
        font_size=10.5)
    add_para(doc, (
        '其中BM25(d|cᵢ)是文档d在查询段cᵢ下的BM25得分，count(d)是检索到文档d的查询段数量，'
        'n是总查询段数。覆盖率奖励因子(1 + count(d)/n)使得被多个独立查询段同时命中的文档'
        '获得额外加分——这是"多源一致性"的强信号，表明该文档可能是可疑文档中多个段落的真实来源，'
        '而非偶发的单一匹配。对于产生单一查询段的短文档，回退为标准BM25检索（max_terms=100）。',
    ))

    # ========== 4. EXPERIMENT ==========
    add_heading_styled(doc, '4 实验', level=1)

    add_heading_styled(doc, '4.1 实验设置', level=2)

    add_heading_styled(doc, '4.1.1 实验数据', level=3)
    add_para(doc, (
        '实验采用PAN26格式的训练数据集，来源于PAN 2025评测数据[1]的转换版本。'
        '数据集包含：源文档语料库60,592篇（学术论文文本，来自ClueWeb09），'
        '可疑文档42,444篇（LLM生成的抄袭案例），以及43,140条查询-文档相关性标注（qrels）。'
        '为高效验证，从42,444条查询中随机采样了两个子集（随机种子20260520）：'
        '800条查询子集和1,200条查询子集。采样覆盖了完整的查询长度分布范围。'
    ))

    add_heading_styled(doc, '4.1.2 数据预处理', level=3)
    add_para(doc, (
        '所有文本进行小写化和tokenization处理（正则表达式[a-z0-9]+提取词元）。'
        '构建倒排索引（inverted index）：以词元为键，以(doc_idx, term_frequency)'
        '二元组列表为值，共包含1,273,004个唯一词元。BM25参数设置为k₁=1.2，b=0.75。'
    ))

    add_heading_styled(doc, '4.1.3 评价指标', level=3)
    add_para(doc, (
        '采用四项标准信息检索评价指标：（1）R@10和R@100（Recall at K）——'
        '正确源文档出现在前K个检索结果中的查询比例；'
        '（2）nDCG@10（Normalized Discounted Cumulative Gain）——'
        '考虑位置权重的归一化折损累积增益；'
        '（3）MRR（Mean Reciprocal Rank）——平均倒数排名，'
        '衡量第一个正确结果排名的倒数的平均值。所有指标数值越高越好。'
    ))

    add_heading_styled(doc, '4.1.4 基线方法', level=3)
    add_para(doc, (
        '基线方法为标准BM25检索[3]：将整篇可疑文档作为单一查询，'
        '选取IDF最高的100个查询词进行检索，返回top-100排序结果。'
        '为公平对比，基线使用与本文方法完全相同的倒排索引和BM25参数。'
    ))

    add_heading_styled(doc, '4.2 实验结果及分析', level=2)

    add_para(doc, '表1和表2分别报告了800条和1,200条查询样本上的实验结果。', bold=True)

    # Table 1
    add_para(doc, '表1. 800条查询样本上的实验结果（submit_pan26.py精确代码）', bold=True, font_size=9, first_line_indent=False)
    table1 = doc.add_table(rows=5, cols=4)
    table1.style = 'Table Grid'
    headers = ['指标', '标准BM25', '分段BM25（本文方法）', 'Δ']
    data_800 = [
        ['R@10', '0.9200', '0.9700', '+5.4%'],
        ['R@100', '0.9700', '0.9912', '+2.2%'],
        ['nDCG@10', '0.8122', '0.9242', '+13.8%'],
        ['MRR', '0.7771', '0.9092', '+17.0%'],
    ]
    for i, h in enumerate(headers):
        table1.rows[0].cells[i].text = h
    for r, row_data in enumerate(data_800):
        for c, val in enumerate(row_data):
            table1.rows[r+1].cells[c].text = val

    add_para(doc, '', font_size=6, first_line_indent=False)  # spacer

    # Table 2
    add_para(doc, '表2. 1200条查询样本上的实验结果', bold=True, font_size=9, first_line_indent=False)
    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    data_1200 = [
        ['R@10', '0.8825', '0.9692', '+9.8%'],
        ['R@100', '0.9517', '0.9908', '+4.1%'],
        ['nDCG@10', '0.7850', '0.9234', '+17.6%'],
        ['MRR', '0.7532', '0.9086', '+20.6%'],
    ]
    for i, h in enumerate(headers):
        table2.rows[0].cells[i].text = h
    for r, row_data in enumerate(data_1200):
        for c, val in enumerate(row_data):
            table2.rows[r+1].cells[c].text = val

    add_para(doc, '', font_size=6, first_line_indent=False)  # spacer

    add_para(doc, (
        '从实验结果可以得出以下分析：'
    ))

    add_para(doc, (
        '（1）分段溯源建模在所有四项指标上一致优于标准BM25。在800条查询上，R@10从0.920提升至0.970，'
        'MRR从0.777提升至0.909；在1200条查询上，R@10从0.883提升至0.969，MRR从0.753提升至0.909。'
        '两个样本上的绝对性能高度一致，说明方法具有良好的稳定性。'
    ))

    add_para(doc, (
        '（2）MRR的提升幅度最大（+17%至+21%），表明覆盖率加权投票机制能有效地将真实源文档'
        '从较低的排名位置"推"到前列。标准BM25可能将真实源文档排在5-10名附近，'
        '而分段投票通过多段一致性信号将其提升至1-3名。'
    ))

    add_para(doc, (
        '（3）R@100趋近理论上限（0.99），说明在该数据集上源检索的召回率已接近饱和，'
        '剩余的错误案例可能需要在文本对齐阶段进行精细化处理。'
    ))

    add_para(doc, (
        '（4）在1200条查询上，标准BM25的R@10（0.883）低于800条查询上的结果（0.920），'
        '而分段方法的R@10在两个样本上几乎一致（0.970 vs 0.969）。'
        '这暗示分段方法对查询分布的变异性具有更强的鲁棒性。'
    ))

    # ========== 5. CONCLUSION ==========
    add_heading_styled(doc, '5 结论', level=1)

    add_para(doc, (
        '本文针对生成式抄袭检测中的源检索任务，提出了一种基于查询分段溯源建模的方法。'
        '该方法的核心思想是将可疑文档的"多源合成"特性显式纳入检索过程：'
        '通过段落级查询分段将可疑文档拆分为独立查询段，每个查询段独立检索候选源文档，'
        '再通过覆盖率加权投票机制聚合各段结果，构建显式的溯源图。'
        '在PAN26评测数据上的实验表明，该方法在R@10、R@100、nDCG@10和MRR四项指标上'
        '均一致优于标准BM25基线，其中MRR的相对提升幅度最高（约+20%）。'
        '该方法实现轻量化（仅依赖numpy），已成功提交至TIRA评测平台。'
    ))

    add_para(doc, (
        '未来工作方向包括：（1）探索基于章节标题（如Introduction、Method、Results）的'
        '语义分段策略以替代纯段落分段；（2）引入稠密检索模型（如E5）增强对高度改写段落的语义匹配能力；'
        '（3）将溯源图输出直接输入文本对齐阶段，实现源检索与抄袭定位的紧密耦合。'
    ))

    # ========== ACKNOWLEDGMENTS ==========
    add_heading_styled(doc, '6 致谢', level=1)

    add_para(doc, (
        'This work is supported by the National Social Science Foundation of China (24BYY080). '
        '感谢PAN@CLEF评测组织者提供的数据集和评测平台。'
    ))

    # ========== REFERENCES ==========
    add_heading_styled(doc, '参考文献', level=1)

    refs = [
        '[1] Greiner-Petter A, Fröbe M, Wahle J P, et al. Overview of the Plagiarism Detection Task at PAN 2025[C]//'
        'Proceedings of the 16th International Conference of the CLEF Association (CLEF 2025). '
        'CEUR Workshop Proceedings, Vol. 4038, 2025.',

        '[2] Potthast M, Gollub T, Hagen M, et al. Overview of the 6th International Competition on Plagiarism Detection[C]//'
        'Working Notes of CLEF 2014 - Conference and Labs of the Evaluation Forum. '
        'CEUR Workshop Proceedings, Vol. 1180, 2014: 975-997.',

        '[3] Robertson S, Zaragoza H. The Probabilistic Relevance Framework: BM25 and Beyond[J]. '
        'Foundations and Trends in Information Retrieval, 2009, 3(4): 333-389.',

        '[4] Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]//'
        'Proceedings of the 2019 Conference on Empirical Methods in Natural Language Processing (EMNLP). '
        'ACL, 2019: 3982-3992.',

        '[5] Wang L, Yang N, Huang X, et al. Text Embeddings by Weakly-Supervised Contrastive Pre-training[C]//'
        'Proceedings of the 2024 Conference on Empirical Methods in Natural Language Processing (EMNLP). '
        'ACL, 2024.',

        '[6] Karpukhin V, Oguz B, Min S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]//'
        'Proceedings of the 2020 Conference on Empirical Methods in Natural Language Processing (EMNLP). '
        'ACL, 2020: 6769-6781.',

        '[7] Bevendorff J, Dementieva D, Fröbe M, et al. Overview of PAN 2025: Generative AI Detection, '
        'Multilingual Text Detoxification, Multi-author Writing Style Analysis, and Generative Plagiarism Detection '
        '- Extended Abstract[C]//Proceedings of the 47th European Conference on Information Retrieval (ECIR 2025). '
        'Springer LNCS, Vol. 15576, 2025: 434-441.',
    ]

    for ref in refs:
        add_para(doc, ref, font_size=9, first_line_indent=False)

    # Save
    output = Path(__file__).parent.parent / 'docs' / 'PAN26_report_v1.docx'
    doc.save(str(output))
    print(f'Report saved to: {output}')


if __name__ == '__main__':
    main()
